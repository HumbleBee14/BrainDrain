"""Document parsing activity — extracts structured text from uploaded files.

Supports PDF (PyMuPDF), DOCX (python-docx), TXT, Markdown, HTML, and CSV.
CPU-only — no GPU dependencies.
"""

import csv
import io
import json
import logging
from dataclasses import dataclass
from typing import Protocol

import fitz  # PyMuPDF
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from langdetect import LangDetectException, detect
from markdown import markdown
from temporalio import activity

from src import clients, s3_paths

logger = logging.getLogger("platform.parse")


@dataclass
class ParseDocumentInput:
    tenant_id: str
    project_id: str
    document_id: str
    storage_path: str
    mime_type: str


@dataclass
class ParseDocumentOutput:
    page_count: int
    language: str | None
    parse_quality: float
    parsed_storage_path: str


# ── Parser Protocol & Registry ──


class DocumentParser(Protocol):
    """Protocol for document parsers. Implement this to add new format support."""

    @property
    def name(self) -> str:
        """Parser name for metadata."""
        ...

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        """Check if this parser handles the given mime type / file extension."""
        ...

    def parse(self, raw_bytes: bytes) -> list[dict]:
        """Parse raw bytes into structured pages."""
        ...


_PARSER_REGISTRY: list[DocumentParser] = []


def register_parser(parser: DocumentParser) -> DocumentParser:
    """Register a parser in the global registry."""
    _PARSER_REGISTRY.append(parser)
    return parser


def get_parser(mime_type: str, storage_path: str) -> DocumentParser:
    """Find the first parser that can handle this file type."""
    mime = mime_type.lower()
    for parser in _PARSER_REGISTRY:
        if parser.can_parse(mime, storage_path):
            return parser
    # Fallback to plain text
    return _PLAIN_TEXT_PARSER


# ── Concrete Parsers ──


class PdfParser:
    """Extract text and structure from PDF using PyMuPDF."""

    @property
    def name(self) -> str:
        return "pymupdf"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type == "application/pdf" or storage_path.endswith(".pdf")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        pages = []
        doc = fitz.open(stream=raw_bytes, filetype="pdf")

        for page_num, page in enumerate(doc, start=1):
            blocks = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)["blocks"]
            sections = []
            page_text_parts = []

            for block in blocks:
                if block["type"] == 0:  # text block
                    for line in block.get("lines", []):
                        text = "".join(span["text"] for span in line.get("spans", []))
                        if not text.strip():
                            continue

                        # Detect headings by font size
                        max_size = max(
                            (span.get("size", 12) for span in line.get("spans", [])),
                            default=12,
                        )
                        section_type = "heading" if max_size > 14 else "paragraph"

                        sections.append({"type": section_type, "content": text.strip()})
                        page_text_parts.append(text.strip())

            pages.append(
                {
                    "page_num": page_num,
                    "text": "\n".join(page_text_parts),
                    "sections": sections,
                }
            )

        doc.close()
        return pages


class DocxParser:
    """Extract text and structure from DOCX using python-docx."""

    @property
    def name(self) -> str:
        return "python-docx"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ) or storage_path.endswith(".docx")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        doc = DocxDocument(io.BytesIO(raw_bytes))
        sections = []
        text_parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level = 1
                try:
                    level = int(style_name.replace("Heading ", "").replace("Heading", "1"))
                except ValueError:
                    pass
                sections.append({"type": "heading", "level": level, "content": text})
            else:
                sections.append({"type": "paragraph", "content": text})
            text_parts.append(text)

        # Extract tables
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            if rows:
                sections.append({"type": "table", "rows": rows})
                text_parts.append(" | ".join(rows[0]) if rows else "")

        # DOCX doesn't have page numbers, treat as single page
        return [{"page_num": 1, "text": "\n".join(text_parts), "sections": sections}]


class HtmlParser:
    """Extract text and structure from HTML using BeautifulSoup."""

    @property
    def name(self) -> str:
        return "beautifulsoup"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type == "text/html" or storage_path.endswith(".html")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        text = raw_bytes.decode("utf-8", errors="replace")
        soup = BeautifulSoup(text, "html.parser")

        # Remove script and style elements
        for tag in soup(["script", "style"]):
            tag.decompose()

        sections = []
        text_parts = []

        for element in soup.find_all(["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "td"]):
            content = element.get_text(strip=True)
            if not content:
                continue

            tag_name = element.name
            if tag_name.startswith("h"):
                level = int(tag_name[1])
                sections.append({"type": "heading", "level": level, "content": content})
            else:
                sections.append({"type": "paragraph", "content": content})
            text_parts.append(content)

        return [{"page_num": 1, "text": "\n".join(text_parts), "sections": sections}]


class MarkdownParser:
    """Parse Markdown by converting to HTML then extracting structure."""

    _html_parser: HtmlParser

    def __init__(self, html_parser: HtmlParser) -> None:
        self._html_parser = html_parser

    @property
    def name(self) -> str:
        return "markdown"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type == "text/markdown" or storage_path.endswith(".md")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        md_text = raw_bytes.decode("utf-8", errors="replace")
        html = markdown(md_text)
        return self._html_parser.parse(html.encode("utf-8"))


class CsvParser:
    """Parse CSV into structured tabular format."""

    @property
    def name(self) -> str:
        return "csv"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type == "text/csv" or storage_path.endswith(".csv")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        text = raw_bytes.decode("utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return [{"page_num": 1, "text": "", "sections": []}]

        headers = rows[0]
        data_rows = rows[1:]

        sections = [{"type": "table", "headers": headers, "rows": data_rows}]
        text_repr = "\n".join(" | ".join(row) for row in rows)
        return [{"page_num": 1, "text": text_repr, "sections": sections}]


class PlainTextParser:
    """Parse plain text — split by paragraphs. Also serves as the fallback parser."""

    @property
    def name(self) -> str:
        return "plaintext"

    def can_parse(self, mime_type: str, storage_path: str) -> bool:
        return mime_type.startswith("text/") or storage_path.endswith(".txt")

    def parse(self, raw_bytes: bytes) -> list[dict]:
        text = raw_bytes.decode("utf-8", errors="replace")
        paragraphs = [p.strip() for p in text.split("\n\n") if p.strip()]

        sections = [{"type": "paragraph", "content": p} for p in paragraphs]
        return [{"page_num": 1, "text": text, "sections": sections}]


# ── Register parsers (order matters — PlainTextParser must be last) ──

_html_parser_instance = HtmlParser()

_PDF_PARSER = register_parser(PdfParser())
_DOCX_PARSER = register_parser(DocxParser())
_HTML_PARSER = register_parser(_html_parser_instance)
_MARKDOWN_PARSER = register_parser(MarkdownParser(_html_parser_instance))
_CSV_PARSER = register_parser(CsvParser())
_PLAIN_TEXT_PARSER = register_parser(PlainTextParser())


# ── Dispatch functions ──


def _parse_by_type(raw_bytes: bytes, mime_type: str, storage_path: str) -> list[dict]:
    """Route to the appropriate parser based on mime type."""
    parser = get_parser(mime_type, storage_path)
    return parser.parse(raw_bytes)


def _parser_name(mime_type: str) -> str:
    """Return parser name for metadata."""
    parser = get_parser(mime_type, "")
    return parser.name


# ── Main activity ──


@activity.defn
async def parse_document(input: ParseDocumentInput) -> ParseDocumentOutput:
    """Parse an uploaded document into structured JSON and store in S3."""
    s3 = clients.get_s3()
    bucket = clients.get_s3_bucket()
    db = await clients.get_db()

    # Idempotency: skip if already parsed
    status = await db.fetchval("SELECT status FROM documents WHERE id = $1", input.document_id)
    if status == "parsed":
        parsed_key = s3_paths.parsed_path(input.tenant_id, input.project_id, input.document_id)
        activity.logger.info("Document %s already parsed, skipping", input.document_id)
        return ParseDocumentOutput(
            page_count=0, language=None, parse_quality=1.0, parsed_storage_path=parsed_key
        )

    # Update status to parsing
    await db.execute(
        "UPDATE documents SET status = 'parsing', updated_at = now() WHERE id = $1",
        input.document_id,
    )

    try:
        # Download raw file from S3
        activity.heartbeat("downloading")
        response = s3.get_object(Bucket=bucket, Key=input.storage_path)
        raw_bytes = response["Body"].read()

        # Route to parser by mime type
        activity.heartbeat("parsing")
        pages = _parse_by_type(raw_bytes, input.mime_type, input.storage_path)

        # Detect language from combined text
        full_text = " ".join(p["text"] for p in pages if p.get("text"))
        language = _detect_language(full_text)

        # Compute quality score
        parse_quality = _compute_quality(pages, len(raw_bytes))

        # Build structured output
        parsed_output = {
            "version": "1.0",
            "doc_id": input.document_id,
            "parser": _parser_name(input.mime_type),
            "page_count": len(pages),
            "language": language,
            "parse_quality": parse_quality,
            "pages": pages,
        }

        # Upload parsed JSON to S3
        activity.heartbeat("uploading_result")
        parsed_key = s3_paths.parsed_path(input.tenant_id, input.project_id, input.document_id)
        parsed_json = json.dumps(parsed_output, ensure_ascii=False)
        s3.put_object(
            Bucket=bucket,
            Key=parsed_key,
            Body=parsed_json.encode("utf-8"),
            ContentType="application/json",
        )

        # Update DB
        await db.execute(
            "UPDATE documents SET status = 'parsed', parse_quality = $2, "
            "page_count = $3, language = $4, updated_at = now() WHERE id = $1",
            input.document_id,
            parse_quality,
            len(pages),
            language,
        )

        activity.logger.info(
            "Parsed document %s: %d pages, quality=%.2f, lang=%s",
            input.document_id,
            len(pages),
            parse_quality,
            language,
        )

        return ParseDocumentOutput(
            page_count=len(pages),
            language=language,
            parse_quality=parse_quality,
            parsed_storage_path=parsed_key,
        )

    except Exception as e:
        # Mark as failed in DB
        await db.execute(
            "UPDATE documents SET status = 'failed', error_message = $2, "
            "updated_at = now() WHERE id = $1",
            input.document_id,
            str(e)[:500],
        )
        raise


# ── Helpers ──


def _detect_language(text: str) -> str | None:
    """Detect language from text, returns ISO 639-1 code."""
    if not text or len(text) < 20:
        return None
    try:
        return detect(text[:5000])
    except LangDetectException:
        return None


def _compute_quality(pages: list[dict], original_size: int) -> float:
    """Heuristic quality score (0.0 to 1.0)."""
    if not pages:
        return 0.0

    total_chars = sum(len(p.get("text", "")) for p in pages)
    if total_chars == 0:
        return 0.0

    # Text density: expect ~300-500 chars per page
    chars_per_page = total_chars / max(len(pages), 1)
    density_score = min(chars_per_page / 400.0, 1.0)

    # Structure: did we detect any headings?
    has_structure = any(s.get("type") == "heading" for p in pages for s in p.get("sections", []))
    structure_score = 1.0 if has_structure else 0.5

    # Encoding: check for replacement characters
    total_text = " ".join(p.get("text", "") for p in pages)
    replacement_count = total_text.count("\ufffd")
    replacement_ratio = replacement_count / max(total_chars, 1)
    encoding_score = max(0.0, 1.0 - replacement_ratio * 10)

    return round((density_score + structure_score + encoding_score) / 3.0, 2)
